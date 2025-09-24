#!/usr/bin/env python3
"""
Simplified script to generate a Word document from MFSR data CSV file.
This version is easier to modify and understand.
"""

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document():
    """Create Word document from CSV data."""
    
    # Read the CSV file
    csv_file = "full_mfsr_data_with_ids.csv"
    df = pd.read_csv(csv_file)
    
    # Remove empty rows
    df = df.dropna(subset=['Sector', 'Project Name'])
    
    # Filter out documents with type 'Štúdia uskutočniteľnosti'
    print("Filtering out 'Štúdia uskutočniteľnosti' documents...")
    original_count = len(df)
    df = df[df['Type'] != 'štúdia uskutočniteľnosti']
    filtered_count = len(df)
    print(f"Excluded {original_count - filtered_count} 'Štúdia uskutočniteľnosti' documents")
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('MFSR Project Documents', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Group by sector
    for sector_name, sector_data in df.groupby('Sector'):
        # Add sector heading
        doc.add_heading(f'Sector: {sector_name}', level=1)
        
        # Group by project within sector
        project_number = 1
        for project_name, project_data in sector_data.groupby('Project Name'):
            # Add project heading
            doc.add_heading(f'Project {project_number}: {project_name}', level=2)
            
            # Add each document for this project
            for idx, row in project_data.iterrows():
                # Get document name from URL
                url = row['URL']
                if pd.isna(url) or url == '':
                    doc_name = f"Document {idx}"
                else:
                    doc_name = url.split('/')[-1]
                    if doc_name == '':
                        doc_name = f"Document {idx}"
                
                # Add document heading
                doc.add_heading(f'Document: {doc_name}', level=3)
                
                # Add hyperlink if URL exists
                if pd.notna(url) and url != '':
                    para = doc.add_paragraph()
                    para.add_run(f"Link: ").bold = True
                    para.add_run(url)
                
                # Add document details
                details = doc.add_paragraph()
                details.add_run("Document Details:\n").bold = True
                details.add_run(f"📄 File Type: ").bold = True
                details.add_run(f"{row['Type'] if pd.notna(row['Type']) else 'N/A'}\n").bold = True
                details.add_run(f"🆔 Document ID: {row['Document_ID'] if pd.notna(row['Document_ID']) else 'N/A'}\n")
                details.add_run(f"📅 Date: {row['Date'] if pd.notna(row['Date']) else 'N/A'}\n")
                details.add_run(f"💾 File Size: {row['File Size'] if pd.notna(row['File Size']) else 'N/A'}\n")
                
                doc.add_paragraph()  # Add spacing
            
            project_number += 1
            doc.add_paragraph()  # Add spacing between projects
        
        doc.add_paragraph()  # Add spacing between sectors
    
    # Save document
    output_file = "MFSR_Project_Documents.docx"
    doc.save(output_file)
    print(f"Word document created: {output_file}")
    
    # Print summary
    print(f"Sectors: {df['Sector'].nunique()}")
    print(f"Projects: {df.groupby(['Sector', 'Project Name']).ngroups}")
    print(f"Documents (excluding 'Štúdia uskutočniteľnosti'): {len(df)}")

if __name__ == "__main__":
    create_word_document()
