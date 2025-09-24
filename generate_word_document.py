#!/usr/bin/env python3
"""
Script to generate a Word document from the MFSR data CSV file.
Organizes data by sector and project, with hyperlinked document names.
"""

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add blue color and underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def generate_word_document(csv_file_path, output_file_path):
    """
    Generate a Word document from the CSV data.
    
    Args:
        csv_file_path (str): Path to the input CSV file
        output_file_path (str): Path for the output Word document
    """
    
    # Read the CSV file
    print(f"Reading CSV file: {csv_file_path}")
    df = pd.read_csv(csv_file_path)
    
    # Remove any empty rows
    df = df.dropna(subset=['Sector', 'Project Name'])
    
    # No filtering - process all documents
    print(f"Processing all {len(df)} documents...")
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('MFSR Project Documents', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Group data by sector
    sectors = df.groupby('Sector')
    
    for sector_name, sector_data in sectors:
        # Add sector heading
        sector_heading = doc.add_heading(f'Sector: {sector_name}', level=1)
        
        # Group projects within each sector
        projects = sector_data.groupby('Project Name')
        
        project_number = 1
        for project_name, project_data in projects:
            # Add project heading
            project_heading = doc.add_heading(f'Project {project_number}: {project_name}', level=2)
            
            # Add documents for this project
            for idx, row in project_data.iterrows():
                # Create document name from URL (extract filename)
                url = row['URL']
                if pd.isna(url) or url == '':
                    doc_name = f"Document {idx}"
                else:
                    # Extract filename from URL
                    doc_name = url.split('/')[-1]
                    if doc_name == '':
                        doc_name = f"Document {idx}"
                
                # Add document heading with hyperlink directly attached
                if pd.notna(url) and url != '':
                    # Create a paragraph for the document name with hyperlink
                    doc_para = doc.add_paragraph()
                    doc_para.add_run("Document: ").bold = True
                    add_hyperlink(doc_para, doc_name, url)
                else:
                    # Fallback if no URL
                    doc_heading = doc.add_heading(f'Document: {doc_name}', level=3)
                
                # Add document details in a clean format
                details_para = doc.add_paragraph()
                
                # Type of Document
                doc_type = row['Type'] if pd.notna(row['Type']) else 'N/A'
                details_para.add_run(f"Type: ").bold = True
                details_para.add_run(f"{doc_type}")
                
                # Document ID
                doc_id = row['Document_ID'] if pd.notna(row['Document_ID']) else 'N/A'
                details_para.add_run(f" | ID: ").bold = True
                details_para.add_run(f"{doc_id}")
                
                # Date
                date = row['Date'] if pd.notna(row['Date']) else 'N/A'
                details_para.add_run(f" | Date: ").bold = True
                details_para.add_run(f"{date}")
                
                # File Size
                file_size = row['File Size'] if pd.notna(row['File Size']) else 'N/A'
                details_para.add_run(f" | Size: ").bold = True
                details_para.add_run(f"{file_size}")
                
                # Add single line spacing
                doc.add_paragraph()
            
            project_number += 1
    
    # Save the document
    print(f"Saving Word document: {output_file_path}")
    doc.save(output_file_path)
    print("Word document generated successfully!")

def main():
    """Main function to run the script."""
    
    # Define file paths
    csv_file_path = "/Users/ghinwamoujaes/Desktop/World Bank/Code/slovakia_CaseStudy/full_mfsr_data_completewith_ids.csv"
    output_file_path = "/Users/ghinwamoujaes/Desktop/World Bank/Code/slovakia_CaseStudy/MFSR_Project_Documents.docx"
    
    # Check if CSV file exists
    if not os.path.exists(csv_file_path):
        print(f"Error: CSV file not found at {csv_file_path}")
        return
    
    try:
        generate_word_document(csv_file_path, output_file_path)
        print(f"\nDocument successfully created at: {output_file_path}")
        
        # Print summary statistics
        df_summary = pd.read_csv(csv_file_path)
        df_summary = df_summary.dropna(subset=['Sector', 'Project Name'])
        
        sectors = df_summary['Sector'].nunique()
        projects = df_summary.groupby(['Sector', 'Project Name']).ngroups
        total_docs = len(df_summary)
        
        print(f"\nSummary:")
        print(f"- Total Sectors: {sectors}")
        print(f"- Total Projects: {projects}")
        print(f"- Total Documents: {total_docs}")
        
    except Exception as e:
        print(f"Error generating document: {str(e)}")

if __name__ == "__main__":
    main()
